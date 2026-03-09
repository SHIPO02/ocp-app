import streamlit as st
import pandas as pd
import re
import os
import io
from docx import Document

st.set_page_config(page_title="OCP - Suivi Production MFS", layout="wide", page_icon="🚢")

st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Barlow:wght@400;600;700&family=Barlow+Condensed:wght@600;700&display=swap');
        :root {
            --ocp-green: #00843D;
            --ocp-dark:  #005C2A;
            --jerf-color: #00843D;
            --safi-color: #1A6FA8;
            --total-color: #C05A00;
        }
        html, body, [class*="css"] { font-family: 'Barlow', sans-serif; }
        .stApp { background-color: #F4F7F5; }
        h1, h2, h3 { color: var(--ocp-dark) !important; font-family: 'Barlow Condensed', sans-serif !important; }
        .kpi-card {
            border-radius: 12px; padding: 20px 24px; color: white;
            box-shadow: 0 4px 16px rgba(0,0,0,0.12); position: relative; overflow: hidden;
        }
        .kpi-card::before {
            content: ''; position: absolute; top: -20px; right: -20px;
            width: 100px; height: 100px; border-radius: 50%; background: rgba(255,255,255,0.1);
        }
        .kpi-card.jerf  { background: linear-gradient(135deg, #00843D, #005C2A); }
        .kpi-card.safi  { background: linear-gradient(135deg, #1A6FA8, #0D4A73); }
        .kpi-card.total { background: linear-gradient(135deg, #C05A00, #8A3F00); }
        .kpi-label { font-size: 12px; font-weight: 600; opacity: 0.85; letter-spacing: 1px; text-transform: uppercase; margin-bottom: 6px; }
        .kpi-value { font-family: 'Barlow Condensed', sans-serif; font-size: 38px; font-weight: 700; line-height: 1; }
        .kpi-sub   { font-size: 11px; opacity: 0.7; margin-top: 4px; }
        .section-header {
            display: flex; align-items: center; gap: 10px;
            padding: 10px 16px; border-radius: 8px; margin: 20px 0 10px 0;
            font-family: 'Barlow Condensed', sans-serif; font-size: 20px; font-weight: 700; color: white;
        }
        .section-header.jerf  { background: var(--jerf-color); }
        .section-header.safi  { background: var(--safi-color); }
        .section-header.total { background: var(--total-color); }
        .stDownloadButton>button {
            background-color: var(--ocp-green) !important; color: white !important;
            border-radius: 8px !important; border: none !important;
            font-family: 'Barlow', sans-serif !important; font-weight: 600 !important;
        }
        [data-testid="stSidebar"] { border-right: 3px solid var(--ocp-green); background: #FAFFF9; }
        hr { border-color: #D0E8D9 !important; }
    </style>
""", unsafe_allow_html=True)

# ─── HELPERS ────────────────────────────────────────────────────────────────

def force_nombre(valeur):
    if pd.isna(valeur): return 0.0
    if isinstance(valeur, (int, float)):
        return 0.0 if abs(valeur) < 1e-6 else float(valeur)
    s = str(valeur).strip()
    if s in ("-", "", "nan"): return 0.0
    nettoye = re.sub(r'[^\d]', '', s.replace("\xa0", "").replace(" ", ""))
    if len(nettoye) > 12: return 0.0
    try:
        return float(nettoye)
    except ValueError:
        return 0.0

def fmt_number(n):
    return f"{int(n):,}".replace(",", " ")

def generate_word(df_result, titre, periode):
    doc = Document()
    doc.add_heading(titre, 0)
    doc.add_paragraph(f"Période : {periode}")
    table = doc.add_table(rows=1, cols=len(df_result.columns))
    table.style = 'Table Grid'
    for i, col in enumerate(df_result.columns):
        table.rows[0].cells[i].text = col
    for _, row in df_result.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = fmt_number(val) if isinstance(val, (int, float)) else str(val)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def export_buttons(df, prefix, titre, periode, key_suffix=""):
    c1, c2 = st.columns(2)
    with c1:
        st.download_button(f"⬇️ CSV — {prefix}", df.to_csv(index=False).encode('utf-8'),
                           file_name=f"{prefix}_{periode}.csv", mime="text/csv",
                           key=f"csv_{key_suffix}")
    with c2:
        st.download_button(f"⬇️ Word — {prefix}", generate_word(df, titre, periode),
                           file_name=f"{prefix}_{periode}.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           key=f"word_{key_suffix}")

def extract_mois_label(date_str):
    """dd/mm/yyyy → 'Jan 2026'"""
    try:
        parts = str(date_str).split("/")
        if len(parts) == 3:
            m = int(parts[1]); y = parts[2]
            noms = {1:"Jan",2:"Fév",3:"Mar",4:"Avr",5:"Mai",6:"Jun",
                    7:"Jul",8:"Aoû",9:"Sep",10:"Oct",11:"Nov",12:"Déc"}
            return f"{noms.get(m,'?')} {y}"
    except:
        pass
    return "Inconnu"

def normalise_sheet_name(name):
    """Normalise un nom de feuille pour la comparaison : minuscules, sans accents simples, sans espaces."""
    n = name.strip().lower()
    replacements = {"é":"e","è":"e","ê":"e","à":"a","â":"a","û":"u","ô":"o","î":"i","ç":"c"}
    for k, v in replacements.items():
        n = n.replace(k, v)
    return n

# Mots-clés qui indiquent une feuille de données mensuelles valides
# On inclut tout sauf les feuilles récap (TOTAL, RECAP, ANNEE, etc.)
SKIP_KEYWORDS = ["total", "recap", "recapitulatif", "annee", "annuel", "bilan", "synthese", "summary"]

def is_data_sheet(name):
    norm = normalise_sheet_name(name)
    for kw in SKIP_KEYWORDS:
        if kw in norm:
            return False
    return True

# ─── HEADER ──────────────────────────────────────────────────────────────────
col_logo, col_title = st.columns([1, 5])
with col_logo:
    if os.path.exists("logo_ocp.png"):
        st.image("logo_ocp.png", width=110)
    else:
        st.markdown("<div style='font-size:34px;font-weight:900;color:#00843D;font-family:Barlow Condensed,sans-serif;'>OCP</div>", unsafe_allow_html=True)
with col_title:
    st.title("Suivi de la Production MFS 26")
    st.markdown("##### Reporting Consolidé — Jerf Lasfar & Safi • JPH 2026")

st.divider()

# ─── SIDEBAR ─────────────────────────────────────────────────────────────────
st.sidebar.header("📂 Chargement des fichiers")
file_jerf = st.sidebar.file_uploader("🏭 Fichier Jerf (Reporting-JPH 2026)", type=["xlsx"], key="jerf")
file_safi = st.sidebar.file_uploader("🏗️ Fichier Safi (Suivi Production MFS 26)", type=["xlsx"], key="safi")

# ─── PARSE JERF ──────────────────────────────────────────────────────────────
jerf_df = None
if file_jerf:
    try:
        df_raw = pd.read_excel(file_jerf, sheet_name='EXPORT', header=None)
        coords = {"ENGRAIS": None, "CAMIONS": None, "VL": None}
        for r in range(len(df_raw)):
            lbl = " ".join(df_raw.iloc[r, 0:3].astype(str)).upper()
            if "EXPORT ENGRAIS" in lbl: coords["ENGRAIS"] = r
            if "EXPORT CAMIONS" in lbl: coords["CAMIONS"] = r
            if "VL CAMIONS"     in lbl: coords["VL"] = r

        ligne_dates = df_raw.iloc[2, :]
        cols_data = [j for j in range(3, len(ligne_dates)) if pd.notna(ligne_dates[j])]
        rows = []
        for j in cols_data:
            dt = ligne_dates[j]
            dl = dt.strftime('%d/%m/%Y') if hasattr(dt, 'strftime') else str(dt).split(" ")[0]
            v1 = force_nombre(df_raw.iloc[coords["ENGRAIS"], j]) if coords["ENGRAIS"] is not None else 0.0
            v2 = force_nombre(df_raw.iloc[coords["CAMIONS"], j]) if coords["CAMIONS"] is not None else 0.0
            v3 = force_nombre(df_raw.iloc[coords["VL"], j])      if coords["VL"] is not None else 0.0
            rows.append({"Date": dl, "Export Engrais": v1, "Export Camions": v2,
                         "VL Camions": v3, "TOTAL Jerf": v1 + v2 + v3})
        jerf_df = pd.DataFrame(rows)
    except Exception as e:
        st.sidebar.error(f"Erreur Jerf : {e}")

# ─── PARSE SAFI ──────────────────────────────────────────────────────────────
# Structure : chaque feuille = 1 mois
#   - Ligne 7 (index 6) = 1er jour du mois
#   - Colonne A (index 0) = numéro du jour
#   - Colonne AF (index 31) = TSP Export
#   - Colonne AG (index 32) = TSP ML

safi_df = None
safi_debug = []

if file_safi:
    try:
        xl = pd.ExcelFile(file_safi)
        all_sheets = xl.sheet_names

        st.sidebar.markdown("**Feuilles détectées (Safi) :**")
        st.sidebar.caption(" | ".join(all_sheets))

        # Colonnes fixes : AF=31, AG=32 (0-indexed)
        COL_JOUR    = 1   # Colonne B
        COL_TSP_EXP = 31  # Colonne AF
        COL_TSP_ML  = 32  # Colonne AG
        START_ROW   = 6   # Ligne 7 → index 6 (le 1er jour)

        def parse_mois_annee(sheet_name):
            """
            Extrait mois et année depuis le nom de la feuille.
            Exemples : 'JANV 2026', 'Fév 2026', 'Mars 2026' → (mois_num, annee)
            """
            mois_map = {
                "jan": 1, "fev": 2, "fév": 2, "mar": 3, "avr": 4,
                "mai": 5, "jun": 6, "jui": 6, "jul": 7, "aou": 8,
                "aoû": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12, "déc": 12
            }
            parts = sheet_name.strip().split()
            mois_num = None
            annee = None
            for p in parts:
                p_low = p.lower()[:3]
                if p_low in mois_map:
                    mois_num = mois_map[p_low]
                try:
                    y = int(p)
                    if 2000 <= y <= 2100:
                        annee = y
                except:
                    pass
            return mois_num, annee

        rows = []
        for sheet in all_sheets:
            if not is_data_sheet(sheet):
                continue

            mois_num, annee = parse_mois_annee(sheet)
            if mois_num is None or annee is None:
                safi_debug.append(f"⚠️ {sheet} : impossible d'extraire mois/année — ignorée")
                continue

            dfs = pd.read_excel(file_safi, sheet_name=sheet, header=None)

            # Chercher colonnes TSP si AF/AG n'existent pas
            tsp_exp_col = COL_TSP_EXP
            tsp_ml_col  = COL_TSP_ML

            if dfs.shape[1] <= COL_TSP_ML:
                # Recherche par en-tête dans les 8 premières lignes
                found_exp = False
                for hrow in range(min(8, len(dfs))):
                    row_vals = [str(v).strip().upper() for v in dfs.iloc[hrow]]
                    for ci, v in enumerate(row_vals):
                        if "TSP" in v and "EXPORT" in v:
                            tsp_exp_col = ci; found_exp = True
                        if "TSP" in v and "ML" in v:
                            tsp_ml_col = ci
                if not found_exp:
                    safi_debug.append(f"⚠️ {sheet} : colonnes TSP non trouvées ({dfs.shape[1]} cols) — ignorée")
                    continue

            nb_lignes = 0
            # Lire directement depuis START_ROW (ligne 7)
            for ri in range(START_ROW, len(dfs)):
                jour_val = dfs.iloc[ri, COL_JOUR]

                # Arrêter si on tombe sur une ligne vide ou de total
                if pd.isna(jour_val):
                    continue
                s = str(jour_val).strip()
                if s in ("", "nan") or any(k in s.upper() for k in ["TOTAL", "CUMUL", "MOYENNE", "MOY"]):
                    continue

                # Le numéro du jour doit être un entier 1-31
                try:
                    jour_num = int(float(s))
                except ValueError:
                    continue
                if jour_num < 1 or jour_num > 31:
                    continue

                tsp_exp = force_nombre(dfs.iloc[ri, tsp_exp_col]) if tsp_exp_col < dfs.shape[1] else 0.0
                tsp_ml  = force_nombre(dfs.iloc[ri, tsp_ml_col])  if tsp_ml_col  < dfs.shape[1] else 0.0

                # Construire la date complète : jour/mois/année
                date_str = f"{jour_num:02d}/{mois_num:02d}/{annee}"

                rows.append({
                    "Mois":       sheet,
                    "Jour":       jour_num,
                    "Date":       date_str,
                    "TSP Export": tsp_exp,
                    "TSP ML":     tsp_ml,
                    "TOTAL Safi": tsp_exp + tsp_ml
                })
                nb_lignes += 1

            safi_debug.append(f"✅ {sheet} ({mois_num:02d}/{annee}) : {nb_lignes} jours lus")

        safi_df = pd.DataFrame(rows) if rows else None

        with st.sidebar.expander("🔍 Détail lecture Safi"):
            for d in safi_debug:
                st.sidebar.caption(d)

    except Exception as e:
        st.sidebar.error(f"Erreur Safi : {e}")

# ─── FILTRES SIDEBAR ─────────────────────────────────────────────────────────
st.sidebar.divider()
st.sidebar.header("🔍 Filtrage")

if jerf_df is not None:
    dates_jerf = ["Toutes"] + list(jerf_df["Date"].unique())
    choix_jerf = st.sidebar.selectbox("📅 Période Jerf", dates_jerf)
else:
    choix_jerf = "Toutes"

if safi_df is not None:
    mois_safi   = ["Tous"] + list(safi_df["Mois"].unique())
    choix_safi  = st.sidebar.selectbox("📅 Mois Safi", mois_safi)
    # Filtre par jour — disponible quel que soit le mois sélectionné
    df_pour_jours = safi_df if choix_safi == "Tous" else safi_df[safi_df["Mois"] == choix_safi]
    jours_dispo = ["Tous"] + [str(j) for j in sorted(df_pour_jours["Jour"].unique().tolist())]
    choix_jour_safi = st.sidebar.selectbox("📆 Jour Safi", jours_dispo, key="jour_safi_select")
else:
    choix_safi      = "Tous"
    choix_jour_safi = "Tous"

# ─── CUMULS ──────────────────────────────────────────────────────────────────
cumul_jerf  = float(jerf_df["TOTAL Jerf"].sum()) if jerf_df is not None else 0.0
cumul_safi  = float(safi_df["TOTAL Safi"].sum()) if safi_df is not None else 0.0
cumul_total = cumul_jerf + cumul_safi

# ─── KPI CARDS ───────────────────────────────────────────────────────────────
st.markdown("### 📌 Cumul à Date — Toute la Période")
c1, c2, c3 = st.columns(3)
with c1:
    sub1 = "Export Engrais + Camions + VL" if jerf_df is not None else "⚠️ Fichier non chargé"
    st.markdown(f"""
        <div class="kpi-card jerf">
            <div class="kpi-label">🏭 Total Jerf Lasfar</div>
            <div class="kpi-value">{fmt_number(cumul_jerf)}</div>
            <div class="kpi-sub">{sub1}</div>
        </div>""", unsafe_allow_html=True)
with c2:
    sub2 = "TSP Export + TSP ML" if safi_df is not None else "⚠️ Fichier non chargé"
    st.markdown(f"""
        <div class="kpi-card safi">
            <div class="kpi-label">🏗️ Total Safi</div>
            <div class="kpi-value">{fmt_number(cumul_safi)}</div>
            <div class="kpi-sub">{sub2}</div>
        </div>""", unsafe_allow_html=True)
with c3:
    st.markdown(f"""
        <div class="kpi-card total">
            <div class="kpi-label">📊 Total Jerf + Safi</div>
            <div class="kpi-value">{fmt_number(cumul_total)}</div>
            <div class="kpi-sub">Consolidé toutes unités</div>
        </div>""", unsafe_allow_html=True)

st.divider()

# ─── SECTION JERF ────────────────────────────────────────────────────────────
st.markdown('<div class="section-header jerf">🏭 Jerf Lasfar — Chargement Export</div>', unsafe_allow_html=True)

if jerf_df is not None:
    show_jerf = jerf_df if choix_jerf == "Toutes" else jerf_df[jerf_df["Date"] == choix_jerf]
    st.dataframe(
        show_jerf, use_container_width=True, hide_index=True,
        height=min(500, 45 + 35 * len(show_jerf)),
        column_config={
            "Date":           st.column_config.TextColumn("Date"),
            "Export Engrais": st.column_config.NumberColumn("Export Engrais", format="%d"),
            "Export Camions": st.column_config.NumberColumn("Export Camions", format="%d"),
            "VL Camions":     st.column_config.NumberColumn("VL Camions",     format="%d"),
            "TOTAL Jerf":     st.column_config.NumberColumn("TOTAL Jerf ✅",  format="%d"),
        }
    )
    export_buttons(show_jerf, "Jerf", "Rapport Jerf Lasfar", choix_jerf, "jerf")
    if choix_jerf == "Toutes" and len(jerf_df) > 1:
        st.line_chart(jerf_df.set_index("Date")["TOTAL Jerf"], color="#00843D")
else:
    st.info("⬅️ Chargez le fichier **Reporting-JPH 2026** dans la barre latérale pour voir les données Jerf.")

st.divider()

# ─── SECTION SAFI ────────────────────────────────────────────────────────────
st.markdown('<div class="section-header safi">🏗️ Safi — TSP Export & TSP ML (par Jour)</div>', unsafe_allow_html=True)

if safi_df is not None:
    # Filtrage mois puis jour
    show_safi = safi_df.copy()
    if choix_safi != "Tous":
        show_safi = show_safi[show_safi["Mois"] == choix_safi]
    if choix_jour_safi != "Tous":
        show_safi = show_safi[show_safi["Jour"] == int(choix_jour_safi)]

    # Tableau affiché
    display_safi = show_safi[["Mois", "Jour", "Date", "TSP Export", "TSP ML", "TOTAL Safi"]].copy()

    # Lignes total par mois si vue globale
    if choix_safi == "Tous":
        totaux_mois = show_safi.groupby("Mois")[["TSP Export", "TSP ML", "TOTAL Safi"]].sum().reset_index()
        totaux_mois.insert(1, "Jour", "—")
        totaux_mois.insert(2, "Date", totaux_mois["Mois"].apply(lambda m: f"TOTAL {m}"))
        display_safi = pd.concat([display_safi, totaux_mois], ignore_index=True)

    # Ligne grand total toujours en bas
    grand_total = pd.DataFrame([{
        "Mois": "TOTAL GENERAL", "Jour": "—", "Date": "—",
        "TSP Export": show_safi["TSP Export"].sum(),
        "TSP ML":     show_safi["TSP ML"].sum(),
        "TOTAL Safi": show_safi["TOTAL Safi"].sum()
    }])
    display_safi = pd.concat([display_safi, grand_total], ignore_index=True)

    st.dataframe(
        display_safi, use_container_width=True, hide_index=True,
        height=min(600, 45 + 35 * len(display_safi)),
        column_config={
            "Mois":       st.column_config.TextColumn("Mois"),
            "Jour":       st.column_config.TextColumn("Jour"),
            "Date":       st.column_config.TextColumn("Date"),
            "TSP Export": st.column_config.NumberColumn("TSP Export", format="%d"),
            "TSP ML":     st.column_config.NumberColumn("TSP ML",     format="%d"),
            "TOTAL Safi": st.column_config.NumberColumn("TOTAL Safi ✅", format="%d"),
        }
    )
    label_export = f"{choix_safi}_J{choix_jour_safi}" if choix_jour_safi != "Tous" else choix_safi
    export_buttons(show_safi[["Mois", "Jour", "Date", "TSP Export", "TSP ML", "TOTAL Safi"]],
                   "Safi", "Rapport Safi TSP", label_export, "safi")

    # Graphique (seulement si plusieurs lignes et pas filtré sur 1 jour)
    if len(show_safi) > 1 and choix_jour_safi == "Tous":
        st.line_chart(show_safi.set_index("Date")[["TSP Export", "TSP ML"]])
else:
    st.info("⬅️ Chargez le fichier **SUIVI DE LA PRODUCTION MFS 26** dans la barre latérale pour voir les données Safi.")

st.divider()

# ─── SECTION TOTAL ───────────────────────────────────────────────────────────
st.markdown('<div class="section-header total">📊 Total Consolidé — Jerf + Safi par Jour</div>', unsafe_allow_html=True)

if jerf_df is not None or safi_df is not None:

    # ── Tableau par jour : fusion Jerf + Safi sur la date dd/mm/yyyy ──────
    if jerf_df is not None:
        j_day = jerf_df[["Date", "TOTAL Jerf"]].copy()
    else:
        j_day = pd.DataFrame(columns=["Date", "TOTAL Jerf"])

    if safi_df is not None:
        s_day = safi_df[["Date", "TOTAL Safi"]].copy()
    else:
        s_day = pd.DataFrame(columns=["Date", "TOTAL Safi"])

    # Merge sur la date — outer pour garder tous les jours
    if not j_day.empty and not s_day.empty:
        day_df = pd.merge(j_day, s_day, on="Date", how="outer").fillna(0)
    elif not j_day.empty:
        day_df = j_day.copy(); day_df["TOTAL Safi"] = 0.0
    else:
        day_df = s_day.copy(); day_df["TOTAL Jerf"] = 0.0

    day_df["TOTAL Jerf+Safi"] = day_df["TOTAL Jerf"] + day_df["TOTAL Safi"]

    # Tri chronologique
    def date_sort_key(d):
        try:
            parts = str(d).split("/")
            return (int(parts[2]), int(parts[1]), int(parts[0]))
        except:
            return (9999, 99, 99)

    day_df["_sort"] = day_df["Date"].apply(date_sort_key)
    day_df = day_df.sort_values("_sort").drop(columns=["_sort"]).reset_index(drop=True)

    # Ligne total général
    total_row = pd.DataFrame([{
        "Date":            "TOTAL GENERAL",
        "TOTAL Jerf":      day_df["TOTAL Jerf"].sum(),
        "TOTAL Safi":      day_df["TOTAL Safi"].sum(),
        "TOTAL Jerf+Safi": day_df["TOTAL Jerf+Safi"].sum()
    }])
    disp_day = pd.concat([day_df, total_row], ignore_index=True)

    st.dataframe(
        disp_day, use_container_width=True, hide_index=True,
        height=min(600, 45 + 35 * len(disp_day)),
        column_config={
            "Date":            st.column_config.TextColumn("Date"),
            "TOTAL Jerf":      st.column_config.NumberColumn("Total Jerf",      format="%d"),
            "TOTAL Safi":      st.column_config.NumberColumn("Total Safi",      format="%d"),
            "TOTAL Jerf+Safi": st.column_config.NumberColumn("Total Jerf+Safi", format="%d"),
        }
    )
    export_buttons(disp_day, "Total_Consolide_Jour", "Rapport Total Jerf+Safi par Jour", "Consolide", "total_jour")

    # ── Histogramme par MOIS (Jan/Fév/Mars uniquement les mois présents) ──
    st.markdown("#### 📊 Evolution mensuelle — Jerf vs Safi")

    # Construire mois à partir des dates du tableau journalier
    chart_df = day_df.copy()
    chart_df["Mois"] = chart_df["Date"].apply(extract_mois_label)
    chart_df = chart_df[chart_df["Mois"] != "Inconnu"]
    mois_tot = chart_df.groupby("Mois")[["TOTAL Jerf", "TOTAL Safi"]].sum().reset_index()

    # Trier les mois chronologiquement
    def mois_sort_key(m):
        noms = {"Jan":1,"Fév":2,"Mar":3,"Avr":4,"Mai":5,"Jun":6,
                "Jul":7,"Aoû":8,"Sep":9,"Oct":10,"Nov":11,"Déc":12}
        try:
            parts = m.split()
            return (int(parts[1]), noms.get(parts[0], 99))
        except:
            return (9999, 99)

    mois_tot["_sort"] = mois_tot["Mois"].apply(mois_sort_key)
    mois_tot = mois_tot.sort_values("_sort").drop(columns=["_sort"]).reset_index(drop=True)
    # Garder uniquement mois avec au moins une valeur non nulle
    mois_tot = mois_tot[(mois_tot["TOTAL Jerf"] > 0) | (mois_tot["TOTAL Safi"] > 0)]

    if len(mois_tot) > 0:
        st.bar_chart(mois_tot.set_index("Mois")[["TOTAL Jerf", "TOTAL Safi"]])
    else:
        st.info("Pas encore de données pour le graphique mensuel.")

else:
    st.info("⬅️ Chargez au moins un fichier pour voir le total consolidé.")







